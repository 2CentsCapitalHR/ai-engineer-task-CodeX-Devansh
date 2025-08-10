import streamlit as st
import os
import json
from docx import Document
import io
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# LangChain & AI Model Imports
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_community.llms import Ollama

# --- Configuration & Constants ---
INDEX_STORE_PATH = "adgm_faiss_index_local"
EMBEDDING_MODEL_NAME = 'all-MiniLM-L6-v2'
OLLAMA_MODEL = 'llama3'

COMPANY_INCORPORATION_CHECKLIST = {
    "Articles of Association": "Articles of Association",
    "Memorandum of Association": "Memorandum of Association",
    "UBO Declaration Form": "UBO Declaration Form",
    "Register of Members and Directors": "Register of Members and Directors"
}

# --- Caching for Models ---
@st.cache_resource
def load_models_and_retriever():
    st.info(f"Loading embedding model (the 'Librarian')...")
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
    
    st.info(f"Loading knowledge base from disk...")
    db = FAISS.load_local(INDEX_STORE_PATH, embeddings, allow_dangerous_deserialization=True)
    retriever = db.as_retriever(search_kwargs={"k": 3})

    st.info(f"Connecting to powerful local AI model '{OLLAMA_MODEL}' (the 'Analyst')...")
    llm = Ollama(model=OLLAMA_MODEL, temperature=0)
    
    st.success("✅ AI Agent is ready!")
    return llm, retriever

# --- Core Application Logic ---
def identify_document_type(text, filename):
    text_lower, filename_lower = text.lower(), filename.lower()
    if "articles of association" in text_lower or "aoa" in filename_lower: return "Articles of Association"
    if "memorandum of association" in text_lower or "moa" in filename_lower: return "Memorandum of Association"
    if "employment" in text_lower: return "Employment Contract"
    if "resolution" in text_lower: return "Shareholder Resolution"
    return "Unknown Document"

def verify_document_checklist(uploaded_doc_types):
    required, uploaded = set(COMPANY_INCORPORATION_CHECKLIST.keys()), set(uploaded_doc_types.values())
    return list(required - uploaded)

# --- THE DEFINITIVE BUG FIX for docx commenting ---
# This is the correct and robust way to add comments in python-docx
def add_comment_to_paragraph(paragraph, comment_text, author):
    # Create a comment object
    comment = paragraph.part.comments_part.add_comment(
        comment_text, author=author, initials=author[0:2]
    )
    # Add a comment start and end range around the paragraph's content
    paragraph.add_comment_start(comment_id=comment.id)
    paragraph.add_comment_end(comment_id=comment.id)


def create_marked_up_docx(original_doc_stream, all_findings):
    doc = Document(original_doc_stream)
    for finding in all_findings:
        for para in doc.paragraphs:
            # Match the problematic clause to find the right paragraph
            if finding["clause_text"].strip() in para.text.strip() and len(para.text.strip()) > 0:
                try:
                    comment_text = f"AI Issue: {finding['issue']}\nCitation: {finding['citation']}\nSuggestion: {finding['suggestion']}"
                    add_comment_to_paragraph(para, comment_text, "ADGM Agent")
                except Exception as e:
                    st.warning(f"Could not add comment for clause '{finding['clause_text'][:50]}...' due to an error: {e}")
                break # Move to next finding once comment is added
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- Streamlit UI ---
st.set_page_config(layout="wide", page_title="ADGM Corporate Agent")
st.title("ADGM Corporate Agent")
st.markdown("A legal assistant that finds issues, cites specific ADGM rules, and adds comments to your document.")

if 'models_loaded' not in st.session_state:
    st.session_state.llm, st.session_state.retriever = load_models_and_retriever()
    st.session_state.models_loaded = True

llm, retriever = st.session_state.llm, st.session_state.retriever

prompt_template = """
You are an expert ADGM compliance officer. Your task is to analyze a legal clause strictly based on the provided ADGM context documents.
**Instructions:**
1.  Carefully read the "Clause to Analyze".
2.  Review the "Context from ADGM Rules" to find the specific rule or section that applies.
3.  If you find an issue, you MUST respond in the following strict format.
4.  If there are no issues, you MUST respond with ONLY the words: "No significant issues found."

**Context from ADGM Rules:**
{context}

**Clause to Analyze:**
{question}

**Your Analysis (Strict Format):**
- **Issue:** [Describe the problem clearly and concisely]
- **Citation:** [Cite the exact ADGM law, regulation, or document name and section number from the context. For example: "Per ADGM Employment Regulations 2019, Section 9(1)" or "As per the ADGM Company Set-up Checklist, Stage Two".]
- **Suggestion:** [Provide a clear, actionable fix for the clause]
"""
PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question"])

qa_chain = RetrievalQA.from_chain_type(
    llm=llm, chain_type="stuff", retriever=retriever,
    chain_type_kwargs={"prompt": PROMPT}, return_source_documents=False
)

uploaded_files = st.file_uploader(
    "1. Upload your .docx documents for 'Company Incorporation'",
    type="docx", accept_multiple_files=True
)

if st.button("🔍 Analyze Documents") and uploaded_files:
    if not os.path.exists(INDEX_STORE_PATH):
        st.error(f"FATAL: Knowledge base not found. Please run ingest_data.py first.")
    else:
        docs_data, doc_types = {}, {}
        for file in uploaded_files:
            # Keep the file object as a byte stream from the start
            file_stream = io.BytesIO(file.getvalue())
            doc = Document(file_stream)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            doc_type = identify_document_type(full_text, file.name)
            docs_data[file.name] = {"content": full_text, "type": doc_type, "file_obj": file}
            doc_types[file.name] = doc_type

        st.header("1. Document Checklist Verification")
        missing_docs = verify_document_checklist(doc_types)
        if not missing_docs:
            st.success("✅ All required documents appear to be present.")
        else:
            st.error(f"❌ Missing Mandatory Documents: **{', '.join(missing_docs)}**")

        st.header("2. Detailed Document Analysis")
        all_findings = []
        for filename, data in docs_data.items():
            st.subheader(f"📄 Reviewing: *{filename}*")
            with st.expander(f"See analysis results for this document", expanded=True):
                paragraphs = [p for p in data['content'].split('\n') if len(p.strip()) > 20]
                progress_bar = st.progress(0, text="Initializing analysis...")
                
                doc_findings = []
                for i, para in enumerate(paragraphs):
                    progress_bar.progress((i + 1) / len(paragraphs), text=f"Analyzing clause {i+1}/{len(paragraphs)}...")
                    response = qa_chain.invoke(para)
                    result = response.get('result', "").strip()

                    if "No significant issues" not in result:
                        issue_match = re.search(r"Issue:\s*(.*)", result, re.IGNORECASE)
                        citation_match = re.search(r"Citation:\s*(.*)", result, re.IGNORECASE)
                        suggestion_match = re.search(r"Suggestion:\s*(.*)", result, re.IGNORECASE)
                        
                        issue = issue_match.group(1).strip().replace("**", "") if issue_match else "Unspecified issue."
                        citation = citation_match.group(1).strip().replace("**", "") if citation_match else "No specific citation found."
                        suggestion = suggestion_match.group(1).strip().replace("**", "") if suggestion_match else "Review required."

                        finding = {"document": filename, "clause_text": para, "issue": issue, "citation": citation, "suggestion": suggestion}
                        doc_findings.append(finding)
                        all_findings.append(finding)
                        
                        st.error(f"**Issue:** {issue}", icon="🚩")
                        st.info(f"**In Clause:** '{para}'", icon="📋")
                        st.info(f"**Citation:** {citation}", icon="⚖️")
                        st.success(f"**Suggestion:** {suggestion}", icon="💡")
                        st.divider()
                
                if not doc_findings:
                    st.success("No major issues found in this document.")

        st.header("3. Final Summary & Downloads")
        json_report = {
            "process": "Company Incorporation", "documents_uploaded": len(uploaded_files),
            "required_documents": len(COMPANY_INCORPORATION_CHECKLIST),
            "missing_documents": missing_docs, "issues_found": [
                {k: v for k, v in f.items() if k != 'clause_text'} for f in all_findings
            ]
        }
        st.json(json_report)
        
        for filename, data in docs_data.items():
            doc_findings_for_file = [f for f in all_findings if f['document'] == filename]
            if doc_findings_for_file:
                # Pass the original byte stream to the function
                original_stream = io.BytesIO(data["file_obj"].getvalue())
                marked_up_stream = create_marked_up_docx(original_stream, doc_findings_for_file)
                st.download_button(
                    label=f"Download Reviewed {filename}", data=marked_up_stream,
                    file_name=f"REVIEWED_{filename}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

elif not uploaded_files:
    st.info("Please upload one or more .docx documents to begin analysis.")